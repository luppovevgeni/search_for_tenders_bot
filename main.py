# -*- coding: utf-8 -*-
from aiogram import Bot, Dispatcher, types, executor
from aiogram.types import ParseMode, InputFile
from aiogram.dispatcher import FSMContext
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.utils.markdown import link
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from datetime import datetime, timedelta
from os import remove
import config, request_zakupki_gov, docx, request_roseltorg, request_etp_ets, request_sber_ast
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from db import users
import keyboards as kb
import states


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run ()
    r._r.append (hyperlink)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    return hyperlink


def site_parse(site, date, title):
    if site == 'ЕИС закупки':
        return request_zakupki_gov.site_request(date, title)
    elif site == 'Росэлторг':
        return request_roseltorg.site_request(date, title)
    elif site == 'Фабрикант':
        return request_etp_ets.site_request(date, title)
    elif site == 'Сбер аст':
        return request_sber_ast.site_request(date, title)


scheduler = AsyncIOScheduler(timezone="Europe/Moscow")
scheduler.start()


bot = Bot(token=config.TOKEN)
Bot.set_current(bot)
storage = MemoryStorage()
dp = Dispatcher(bot, storage=storage)

chars = ['░', '▒', '▓', '█']

@dp.message_handler(commands=['start','home'], state='*')
async def start(message: types.Message, state: FSMContext):
    await state.finish()
    if users.count_documents({'user_id': message.from_user.id}) == 0:
        await message.answer('Для начала пользования ботом, вам надо заполнить *ключевые слова* для поиска, для ежедневной выписки надо заполнить *время рассылки*', parse_mode=ParseMode.MARKDOWN)
        users.insert_one({'user_id': message.from_user.id,
                          'time': None,
                          'keywords': [],
                          'filters' : [],
                          'sites': [],
                          'days': ['0', '1', '2', '3', '4']})
    await message.answer('Личный кабинет', reply_markup=kb.personal_area, parse_mode=ParseMode.MARKDOWN)


@dp.callback_query_handler(state=None)
async def callback(callback_query: types.CallbackQuery, state: FSMContext):
    data = callback_query.data
    if data == 'in_developing':
        await callback_query.answer('В разроботке', show_alert=True)
    elif data == 'posting':
        time = users.find_one({'user_id': callback_query.from_user.id})['time']
        if time == None:
            str_time = 'Рассылка не подключена'
        else:
            str_time = f'Время рассылки: {time}'
        await callback_query.message.delete()
        await callback_query.message.answer(str_time, reply_markup=kb.time_settings)
    elif data == 'keywords':
        kaywords = "\n".join(users.find_one({"user_id": callback_query.from_user.id})["keywords"])
        await callback_query.message.delete()
        await callback_query.message.answer(f'Ключевые слова:\n{kaywords}', reply_markup=kb.keywords_settings)
    elif data == 'site':
        sites = users.find_one({'user_id': callback_query.from_user.id})['sites']
        await callback_query.message.delete()
        await callback_query.message.answer(f'Доступные сайты:', reply_markup=kb.site(sites))
    elif data.split('_')[0] == 'site':
        sites = users.find_one({'user_id': callback_query.from_user.id})['sites']
        if data.split('_')[1] in sites:
            sites.remove(data.split('_')[1])
        else:
            sites.append(data.split('_')[1])
        await callback_query.message.edit_reply_markup(reply_markup=kb.site(sites))
        users.update_one({'user_id': callback_query.from_user.id}, {'$set': {'sites': sites}})
    elif data == 'home':
        await callback_query.message.delete()
        await callback_query.message.answer('Личный кабинет', reply_markup=kb.personal_area)
    elif data == 'change_keywords':
        await states.keywords.first()
        await callback_query.message.delete()
        await callback_query.message.answer('Напиши ключевые слова через запятую, например:\n\nВыполнены ли работы, Техническое обследование')
    elif data == 'change_time':
        await states.posting_time.first()
        await callback_query.message.delete()
        await callback_query.message.answer('Напиши время в которое бот будет посылать информацию каждый день (напрмер 11)')
    elif data == 'del_time':
        try:
            await callback_query.message.delete()
            await callback_query.message.answer('Расылка не подключена', reply_markup=kb.time_settings)
            users.update_one({'user_id': callback_query.from_user.id}, {'$set': {'time': None}})
            scheduler.remove_job(str(callback_query.from_user.id))
            await callback_query.answer('Вы отключили рассылку', show_alert=False)
        except:
            pass
    elif data == 'week_days':
        await callback_query.message.delete()
        days = users.find_one({'user_id': callback_query.from_user.id})['days']
        await callback_query.message.answer('Дни недели по которым будет приходить рассылка', reply_markup=kb.day(days))
    elif data.split('_')[0] == 'day':
        days = users.find_one({'user_id': callback_query.from_user.id})['days']
        if data.split('_')[1] in days:
            days.remove(data.split('_')[1])
        else:
            days.append(data.split('_')[1])
        await callback_query.message.edit_reply_markup(reply_markup=kb.day(days))
        users.update_one({'user_id': callback_query.from_user.id}, {'$set': {'days': days}})



@dp.message_handler(state=states.posting_time.time)
async def time(message: types.Message, state: FSMContext):
    try:
        time = int(message.text)
        users.update_one({'user_id': message.from_user.id}, {'$set': {'time': time}})
        try:
            scheduler.remove_job(str(message.from_user.id))
        except:
            pass
        scheduler.add_job(parse, id=str(message.from_user.id), trigger='cron', hour=time, minute='00', args=[message.from_user.id])
        str_time = f'Время рассылки: {time}'
        await message.answer(str_time, reply_markup=kb.time_settings)
        await state.finish()
    except:
        await message.answer('Ошибка, проверьте время и напишите еще раз')


@dp.message_handler(state=states.keywords.keywords)
async def change_keywords(message: types.Message, state: FSMContext):
    try:
        keywords = list(map(lambda x: x.rstrip().lstrip(), message.text.split(',')))
        users.update_one({'user_id': message.from_user.id}, {'$set': {'keywords': keywords}})
        keywords = "\n".join(keywords)
        await message.answer(f'Ключевые слова:\n{keywords}', reply_markup=kb.keywords_settings)
        await state.finish()
    except:
        await message.answer('Ошибка, проверьте ключевые слова и напишите еще раз')


@dp.message_handler(commands=['parse'], state=None)
async def forcibly_parse(message: types.Message):
    if users.count_documents({'user_id': message.from_user.id}) != 0:
        await parse(message.from_user.id, True)


async def parse(id, forcibly=False):
    date = datetime.now()
    if str(date.weekday()) in users.find_one({'user_id': id})['days'] or forcibly:
        f = False
        request_template = users.find_one({'user_id': id})['keywords']
        doc = docx.Document()
        par = doc.add_paragraph('Выписка по актуальным тендерам')
        await bot.send_message(id, 'Начинаю проверку')
        if request_template == []:
            await bot.send_message(id, 'Заполните *ключевые слова* (/home)', parse_mode=ParseMode.MARKDOWN)
        else:
            sites = users.find_one({'user_id': id})['sites']
            for site in sites:
                message = await bot.send_message(id, f'Проверяю {site}')
                try:
                    for elem in request_template:
                        r = site_parse(site, date, elem)
                        percent = f'{round(((request_template.index(elem) + 1) / len(request_template)) * 100)} %'
                        percent_for_element = (1 / len(request_template)) * 30
                        progress_bar = '█' * int((percent_for_element * (request_template.index(elem) + 1)) // 3) + chars[int((percent_for_element * (request_template.index(elem) + 1)) % 3)]
                        progress_bar = progress_bar + '░' * (10 - len(progress_bar))
                        for text in r:
                            f = True
                            par = doc.add_paragraph('\n'.join(text.split('\n')[:-2]))
                            add_hyperlink(par, '\nПодробнее', text.split('\n')[-2])
                        await message.edit_text(f'Проверяю {site}\n{progress_bar} {percent}')
                    await message.edit_text(f'Проверил {site}')
                except:
                    await message.edit_text(f'Ошибка {site}')
            if not f:
                await bot.send_message(id, 'Нет новых тендеров')
            else:
                doc.save(f'{id}.docx')
                await bot.send_message(id, 'Проверку закончил')
                await bot.send_document(id, InputFile(f'{id}.docx', 'Актуальные тендеры.docx'))
                remove(f'{id}.docx')


if __name__ == "__main__":
    for user in users.find({}):
        if user['time'] != None:
            scheduler.add_job(parse, id=str(user['user_id']) ,trigger='cron', hour=user['time'], minute='00', args=[user['user_id']])
    executor.start_polling(dp, skip_updates=True)